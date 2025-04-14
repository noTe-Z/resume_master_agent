"""
DOCX Resume Parser

This module provides functions for extracting structured data from resume DOCX files.
"""

import logging
from typing import Dict, Any
import docx

# Initialize logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text from the DOCX
    """
    try:
        # Open the DOCX file
        doc = docx.Document(docx_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text:
                            full_text.append(para.text)
        
        # Join the text with newlines
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return ""

def parse_resume_docx(docx_path: str) -> Dict[str, Any]:
    """
    Parse a resume DOCX file into structured data.
    
    This function extracts text from the DOCX and defers to the main parser
    for further processing.
    
    Args:
        docx_path: Path to the resume DOCX file
        
    Returns:
        Dictionary containing structured resume data
    """
    try:
        # Import here to avoid circular imports
        from backend.resume_parser.parser import (
            normalize_text, identify_sections, extract_contact_info,
            parse_experience_section, parse_education_section,
            parse_skills_section, extract_certifications
        )
        
        # Extract text from DOCX
        raw_text = extract_text_from_docx(docx_path)
        if not raw_text:
            logger.error(f"Failed to extract text from {docx_path}")
            return {'error': 'Failed to extract text from DOCX'}
        
        # Normalize the text
        normalized_text = normalize_text(raw_text)
        
        # Identify sections in the resume
        sections = identify_sections(normalized_text)
        
        # Extract contact information (typically from the header section)
        contact_info = extract_contact_info(sections.get('other', '') or normalized_text[:500])
        
        # Parse experience section
        experiences = parse_experience_section(sections.get('experience', ''))
        
        # Parse education section
        education = parse_education_section(sections.get('education', ''))
        
        # Parse skills section
        skills = parse_skills_section(sections.get('skills', ''))
        
        # Extract certifications
        certifications = extract_certifications(sections.get('certifications', ''))
        
        # Create structured resume data
        resume_data = {
            'contact_info': contact_info,
            'summary': sections.get('summary', ''),
            'experiences': experiences,
            'education': education,
            'skills': skills,
            'certifications': certifications,
            'raw_sections': sections  # Include raw sections for debugging
        }
        
        return resume_data
        
    except Exception as e:
        logger.error(f"Error parsing resume DOCX: {e}")
        return {'error': str(e)} 